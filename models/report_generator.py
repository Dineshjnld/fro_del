"""
Report Generator using Pegasus for summaries and visualization
"""
import torch
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import plotly.express as px
import seaborn as sns
import pandas as pd
from jinja2 import Template
import weasyprint
from docx import Document
from pathlib import Path
import logging
import asyncio
from typing import Dict, List, Any, Optional
from datetime import datetime
from transformers import PegasusTokenizer, PegasusForConditionalGeneration
from config.settings import settings

class ReportGenerator:
    """Generate comprehensive reports with AI summaries and visualizations"""
    
    def __init__(self, config: dict):
        self.logger = logging.getLogger(__name__)
        self.config = config
        self.device = torch.device("cuda" if torch.cuda.is_available() and settings.USE_GPU else "cpu")
        
        # Load Pegasus for summarization
        self._load_summarization_model()
        
        # Setup report directories
        self.reports_dir = Path(settings.REPORTS_DIR)
        self.reports_dir.mkdir(exist_ok=True)
        
        # Chart styling
        plt.style.use('seaborn-v0_8')
        sns.set_palette("husl")
    
    def _load_summarization_model(self):
        """Load Pegasus model for text summarization"""
        try:
            model_name = self.config.get("model", "google/pegasus-cnn_dailymail")
            
            self.summary_tokenizer = PegasusTokenizer.from_pretrained(
                model_name,
                cache_dir=settings.MODELS_DIR
            )
            self.summary_model = PegasusForConditionalGeneration.from_pretrained(
                model_name,
                cache_dir=settings.MODELS_DIR
            )
            
            self.summary_model.to(self.device)
            self.summary_model.eval()
            
            self.logger.info(f"✅ Pegasus summarization model loaded on {self.device}")
            
        except Exception as e:
            self.logger.warning(f"⚠️ Failed to load Pegasus: {e}")
            self.summary_model = None
            self.summary_tokenizer = None
    
    async def generate_comprehensive_report(
        self, 
        query_data: Dict[str, Any], 
        results: List[Dict[str, Any]],
        report_type: str = "standard"
    ) -> Dict[str, Any]:
        """
        Generate comprehensive report with summary, charts, and exports
        
        Args:
            query_data: Original query information
            results: Query results data
            report_type: Type of report (standard, executive, detailed)
            
        Returns:
            Dict with report paths and metadata
        """
        try:
            report_id = f"cctns_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            # Generate summary
            summary = await self._generate_ai_summary(query_data, results)
            
            # Generate visualizations
            chart_paths = await self._generate_visualizations(results, report_id)
            
            # Generate reports in multiple formats
            reports = {}
            
            if report_type in ["standard", "executive", "detailed"]:
                reports["html"] = await self._generate_html_report(
                    query_data, results, summary, chart_paths, report_id, report_type
                )
                reports["pdf"] = await self._generate_pdf_report(reports["html"], report_id)
            
            if report_type in ["detailed"]:
                reports["docx"] = await self._generate_docx_report(
                    query_data, results, summary, report_id
                )
                reports["excel"] = await self._generate_excel_report(results, report_id)
            
            # Generate metadata
            metadata = {
                "report_id": report_id,
                "generated_at": datetime.now().isoformat(),
                "query": query_data.get("original_query", ""),
                "sql": query_data.get("sql", ""),
                "result_count": len(results),
                "report_type": report_type,
                "files_generated": list(reports.keys()),
                "chart_count": len(chart_paths)
            }
            
            return {
                "success": True,
                "report_id": report_id,
                "summary": summary,
                "charts": chart_paths,
                "reports": reports,
                "metadata": metadata
            }
            
        except Exception as e:
            self.logger.error(f"Report generation failed: {e}")
            return {
                "success": False,
                "error": str(e),
                "summary": "Report generation failed",
                "charts": [],
                "reports": {}
            }
    
    async def _generate_ai_summary(self, query_data: Dict, results: List[Dict]) -> str:
        """Generate AI-powered summary using Pegasus"""
        if not self.summary_model or not results:
            return self._generate_basic_summary(query_data, results)
        
        try:
            # Prepare content for summarization
            content = self._prepare_content_for_summary(query_data, results)
            
            inputs = self.summary_tokenizer(
                content,
                return_tensors="pt",
                max_length=1024,
                truncation=True,
                padding=True
            ).to(self.device)
            
            with torch.no_grad():
                summary_ids = self.summary_model.generate(
                    **inputs,
                    max_length=self.config.get("max_length", 150),
                    min_length=self.config.get("min_length", 30),
                    num_beams=self.config.get("num_beams", 4),
                    early_stopping=True,
                    no_repeat_ngram_size=2
                )
            
            ai_summary = self.summary_tokenizer.decode(summary_ids[0], skip_special_tokens=True)
            
            # Enhance with domain context
            enhanced_summary = self._enhance_summary_with_context(ai_summary, query_data, results)
            
            return enhanced_summary
            
        except Exception as e:
            self.logger.warning(f"AI summary generation failed: {e}")
            return self._generate_basic_summary(query_data, results)
    
    def _prepare_content_for_summary(self, query_data: Dict, results: List[Dict]) -> str:
        """Prepare content for AI summarization"""
        content = f"Police Query Analysis Report\n\n"
        content += f"Original Query: {query_data.get('original_query', 'N/A')}\n"
        content += f"Generated SQL: {query_data.get('sql', 'N/A')}\n"
        content += f"Results Found: {len(results)} records\n\n"
        
        if results:
            content += "Key Findings:\n"
            
            # Sample results for context
            sample_size = min(5, len(results))
            for i, result in enumerate(results[:sample_size]):
                content += f"Record {i+1}: "
                key_values = []
                for key, value in result.items():
                    if value is not None:
                        key_values.append(f"{key}: {value}")
                content += ", ".join(key_values[:3]) + "\n"
            
            # Statistical summary
            if len(results) > 1:
                content += f"\nStatistical Overview:\n"
                # Analyze numeric columns
                df = pd.DataFrame(results)
                numeric_cols = df.select_dtypes(include=['number']).columns
                
                for col in numeric_cols:
                    values = df[col].dropna()
                    if len(values) > 0:
                        content += f"{col}: Total={values.sum()}, Average={values.mean():.2f}, Max={values.max()}, Min={values.min()}\n"
        
        return content
    
    def _generate_basic_summary(self, query_data: Dict, results: List[Dict]) -> str:
        """Generate basic summary when AI model is not available"""
        summary = f"CCTNS Query Report Summary\n\n"
        summary += f"Query: {query_data.get('original_query', 'N/A')}\n"
        summary += f"Results: Found {len(results)} records\n"
        
        if results:
            # Basic analysis
            summary += f"Data Overview:\n"
            if len(results) > 0:
                columns = list(results[0].keys())
                summary += f"- Columns: {', '.join(columns)}\n"
                
                # Count non-null values
                df = pd.DataFrame(results)
                for col in columns[:3]:  # Limit to first 3 columns
                    non_null_count = df[col].notna().sum()
                    summary += f"- {col}: {non_null_count} valid entries\n"
        
        return summary
    
    def _enhance_summary_with_context(self, ai_summary: str, query_data: Dict, results: List[Dict]) -> str:
        """Enhance AI summary with police domain context"""
        enhanced = f"🛡️ CCTNS Database Analysis Report\n\n"
        enhanced += f"📋 Query: {query_data.get('original_query', 'N/A')}\n"
        enhanced += f"📊 Results: {len(results)} records found\n\n"
        enhanced += f"🤖 AI Analysis:\n{ai_summary}\n\n"
        
        # Add domain-specific insights
        if results:
            enhanced += f"🔍 Key Insights:\n"
            
            # Crime analysis
            if any('crime' in str(k).lower() for result in results for k in result.keys()):
                enhanced += "- Crime-related data analysis completed\n"
            
            # Officer analysis  
            if any('officer' in str(k).lower() for result in results for k in result.keys()):
                enhanced += "- Officer performance metrics included\n"
            
            # District analysis
            if any('district' in str(k).lower() for result in results for k in result.keys()):
                enhanced += "- District-wise breakdown available\n"
        
        enhanced += f"\n⏰ Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        return enhanced
    
    async def _generate_visualizations(self, results: List[Dict], report_id: str) -> List[str]:
        """Generate various visualizations for the data"""
        if not results:
            return []
        
        chart_paths = []
        df = pd.DataFrame(results)
        
        try:
            # 1. Bar Chart for categorical data
            categorical_cols = df.select_dtypes(include=['object']).columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            
            if len(categorical_cols) > 0 and len(numeric_cols) > 0:
                chart_path = await self._create_bar_chart(df, categorical_cols[0], numeric_cols[0], report_id)
                if chart_path:
                    chart_paths.append(chart_path)
            
            # 2. Time series chart if date columns exist
            date_cols = [col for col in df.columns if 'date' in col.lower() or 'time' in col.lower()]
            if date_cols and len(numeric_cols) > 0:
                chart_path = await self._create_time_series_chart(df, date_cols[0], numeric_cols[0], report_id)
                if chart_path:
                    chart_paths.append(chart_path)
            
            # 3. Pie chart for distribution
            if len(categorical_cols) > 0:
                chart_path = await self._create_pie_chart(df, categorical_cols[0], report_id)
                if chart_path:
                    chart_paths.append(chart_path)
            
            # 4. Summary statistics chart
            if len(numeric_cols) > 1:
                chart_path = await self._create_summary_chart(df, numeric_cols, report_id)
                if chart_path:
                    chart_paths.append(chart_path)
                    
        except Exception as e:
            self.logger.warning(f"Chart generation error: {e}")
        
        return chart_paths
    
    async def _create_bar_chart(self, df: pd.DataFrame, x_col: str, y_col: str, report_id: str) -> Optional[str]:
        """Create bar chart"""
        try:
            plt.figure(figsize=(12, 6))
            
            # Group and aggregate data
            grouped = df.groupby(x_col)[y_col].sum().sort_values(ascending=False)
            
            # Limit to top 15 items for readability
            if len(grouped) > 15:
                grouped = grouped.head(15)
            
            bars = plt.bar(range(len(grouped)), grouped.values, color='#1565C0')
            plt.xlabel(x_col.replace('_', ' ').title())
            plt.ylabel(y_col.replace('_', ' ').title())
            plt.title(f'{y_col.replace("_", " ").title()} by {x_col.replace("_", " ").title()}')
            plt.xticks(range(len(grouped)), grouped.index, rotation=45, ha='right')
            
            # Add value labels on bars
            for bar, value in zip(bars, grouped.values):
                plt.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.01 * max(grouped.values),
                        f'{value}', ha='center', va='bottom')
            
            plt.tight_layout()
            
            chart_path = str(self.reports_dir / f"{report_id}_bar_chart.png")
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            return chart_path
            
        except Exception as e:
            self.logger.warning(f"Bar chart creation failed: {e}")
            return None
    
    async def _create_time_series_chart(self, df: pd.DataFrame, date_col: str, value_col: str, report_id: str) -> Optional[str]:
        """Create time series chart"""
        try:
            plt.figure(figsize=(12, 6))
            
            # Convert to datetime and sort
            df_copy = df.copy()
            df_copy[date_col] = pd.to_datetime(df_copy[date_col], errors='coerce')
            df_copy = df_copy.dropna(subset=[date_col]).sort_values(date_col)
            
            # Group by date and sum values
            daily_data = df_copy.groupby(df_copy[date_col].dt.date)[value_col].sum()
            
            plt.plot(daily_data.index, daily_data.values, marker='o', linewidth=2, markersize=4, color='#1565C0')
            plt.xlabel('Date')
            plt.ylabel(value_col.replace('_', ' ').title())
            plt.title(f'{value_col.replace("_", " ").title()} Over Time')
            plt.xticks(rotation=45)
            plt.grid(True, alpha=0.3)
            
            plt.tight_layout()
            
            chart_path = str(self.reports_dir / f"{report_id}_time_series.png")
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            return chart_path
            
        except Exception as e:
            self.logger.warning(f"Time series chart creation failed: {e}")
            return None
    
    async def _create_pie_chart(self, df: pd.DataFrame, category_col: str, report_id: str) -> Optional[str]:
        """Create pie chart for distribution"""
        try:
            plt.figure(figsize=(10, 8))
            
            # Count values and limit to top 10
            value_counts = df[category_col].value_counts().head(10)
            
            # Create pie chart
            colors = plt.cm.Set3(range(len(value_counts)))
            wedges, texts, autotexts = plt.pie(
                value_counts.values, 
                labels=value_counts.index, 
                autopct='%1.1f%%',
                colors=colors,
                startangle=90
            )
            
            plt.title(f'Distribution of {category_col.replace("_", " ").title()}')
            
            # Improve text readability
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
            
            plt.axis('equal')
            
            chart_path = str(self.reports_dir / f"{report_id}_pie_chart.png")
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            return chart_path
            
        except Exception as e:
            self.logger.warning(f"Pie chart creation failed: {e}")
            return None
    
    async def _create_summary_chart(self, df: pd.DataFrame, numeric_cols: List[str], report_id: str) -> Optional[str]:
        """Create summary statistics chart"""
        try:
            fig, axes = plt.subplots(2, 2, figsize=(15, 10))
            axes = axes.flatten()
            
            # Select up to 4 numeric columns
            cols_to_plot = numeric_cols[:4]
            
            for i, col in enumerate(cols_to_plot):
                if i >= len(axes):
                    break
                    
                # Create histogram
                axes[i].hist(df[col].dropna(), bins=20, color='#1565C0', alpha=0.7, edgecolor='black')
                axes[i].set_title(f'Distribution of {col.replace("_", " ").title()}')
                axes[i].set_xlabel(col.replace('_', ' ').title())
                axes[i].set_ylabel('Frequency')
                axes[i].grid(True, alpha=0.3)
            
            # Hide unused subplots
            for i in range(len(cols_to_plot), len(axes)):
                axes[i].set_visible(False)
            
            plt.suptitle('Statistical Summary of Numeric Data', fontsize=16, fontweight='bold')
            plt.tight_layout()
            
            chart_path = str(self.reports_dir / f"{report_id}_summary_stats.png")
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            return chart_path
            
        except Exception as e:
            self.logger.warning(f"Summary chart creation failed: {e}")
            return None
    
    async def _generate_html_report(
        self, 
        query_data: Dict, 
        results: List[Dict], 
        summary: str, 
        chart_paths: List[str], 
        report_id: str,
        report_type: str
    ) -> str:
        """Generate HTML report"""
        try:
            template_str = self._get_html_template(report_type)
            template = Template(template_str)
            
            # Prepare data for template
            template_data = {
                "title": "CCTNS Database Analysis Report",
                "report_id": report_id,
                "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "query": query_data.get("original_query", "N/A"),
                "sql": query_data.get("sql", "N/A"),
                "summary": summary.replace('\n', '<br>'),
                "results": results[:100],  # Limit for HTML display
                "total_results": len(results),
                "charts": [Path(chart).name for chart in chart_paths],
                "report_type": report_type.title()
            }
            
            html_content = template.render(**template_data)
            
            html_path = str(self.reports_dir / f"{report_id}.html")
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return html_path
            
        except Exception as e:
            self.logger.error(f"HTML report generation failed: {e}")
            return ""
    
    def _get_html_template(self, report_type: str) -> str:
        """Get HTML template based on report type"""
        return """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ title }}</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; background-color: #f5f5f5; }
        .container { max-width: 1200px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 0 20px rgba(0,0,0,0.1); }
        .header { text-align: center; border-bottom: 3px solid #1565C0; padding-bottom: 20px; margin-bottom: 30px; }
        .header h1 { color: #1565C0; margin: 0; font-size: 2.5em; }
        .header .subtitle { color: #666; font-size: 1.2em; margin-top: 10px; }
        .section { margin: 30px 0; }
        .section h2 { color: #1565C0; border-left: 4px solid #1565C0; padding-left: 15px; }
        .query-box { background: #f8f9fa; border: 1px solid #e9ecef; border-radius: 5px; padding: 20px; font-family: monospace; }
        .sql-box { background: #2d3748; color: #e2e8f0; border-radius: 5px; padding: 20px; font-family: monospace; overflow-x: auto; }
        .summary-box { background: #e8f5e8; border: 1px solid #4CAF50; border-radius: 5px; padding: 20px; }
        .results-table { width: 100%; border-collapse: collapse; margin-top: 20px; }
        .results-table th { background: #1565C0; color: white; padding: 12px; text-align: left; }
        .results-table td { padding: 10px; border-bottom: 1px solid #ddd; }
        .results-table tr:nth-child(even) { background: #f8f9fa; }
        .charts { display: grid; grid-template-columns: repeat(auto-fit, minmax(400px, 1fr)); gap: 20px; margin: 20px 0; }
        .chart-container { text-align: center; background: #f8f9fa; padding: 20px; border-radius: 10px; }
        .chart-container img { max-width: 100%; height: auto; border-radius: 5px; }
        .metadata { background: #f1f3f4; padding: 15px; border-radius: 5px; font-size: 0.9em; color: #555; }
        .footer { text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; color: #666; }
        @media print { .container { box-shadow: none; } }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🛡️ {{ title }}</h1>
            <div class="subtitle">{{ report_type }} Report</div>
            <div class="subtitle">Report ID: {{ report_id }}</div>
            <div class="subtitle">Generated: {{ timestamp }}</div>
        </div>
        
        <div class="section">
            <h2>📋 Query Information</h2>
            <div class="query-box">
                <strong>Original Query:</strong><br>
                {{ query }}
            </div>
            {% if sql %}
            <div class="sql-box">
                <strong>Generated SQL:</strong><br>
                {{ sql }}
            </div>
            {% endif %}
        </div>
        
        <div class="section">
            <h2>🤖 AI-Generated Summary</h2>
            <div class="summary-box">
                {{ summary|safe }}
            </div>
        </div>
        
        {% if charts %}
        <div class="section">
            <h2>📊 Data Visualizations</h2>
            <div class="charts">
                {% for chart in charts %}
                <div class="chart-container">
                    <img src="{{ chart }}" alt="Chart {{ loop.index }}">
                    <p><strong>Chart {{ loop.index }}</strong></p>
                </div>
                {% endfor %}
            </div>
        </div>
        {% endif %}
        
        {% if results %}
        <div class="section">
            <h2>📋 Query Results</h2>
            <p><strong>Showing {{ results|length }} of {{ total_results }} records</strong></p>
            <table class="results-table">
                <thead>
                    <tr>
                        {% for key in results[0].keys() %}
                        <th>{{ key.replace('_', ' ').title() }}</th>
                        {% endfor %}
                    </tr>
                </thead>
                <tbody>
                    {% for row in results %}
                    <tr>
                        {% for value in row.values() %}
                        <td>{{ value if value is not none else 'N/A' }}</td>
                        {% endfor %}
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
        </div>
        {% endif %}
        
        <div class="section">
            <h2>ℹ️ Report Metadata</h2>
            <div class="metadata">
                <strong>Report Type:</strong> {{ report_type }}<br>
                <strong>Total Records:</strong> {{ total_results }}<br>
                <strong>Charts Generated:</strong> {{ charts|length }}<br>
                <strong>Generation Time:</strong> {{ timestamp }}<br>
                <strong>Generated By:</strong> CCTNS Copilot Engine
            </div>
        </div>
        
        <div class="footer">
            <p>📄 Generated by CCTNS Copilot Engine | Andhra Pradesh Police Department</p>
            <p>🤖 Powered by AI4Bharat, OpenAI Whisper, Google FLAN-T5, Microsoft CodeT5, Google Pegasus</p>
        </div>
    </div>
</body>
</html>
        """
    
    async def _generate_pdf_report(self, html_path: str, report_id: str) -> str:
        """Generate PDF from HTML report"""
        try:
            if not html_path or not Path(html_path).exists():
                return ""
            
            pdf_path = str(self.reports_dir / f"{report_id}.pdf")
            
            # Use weasyprint to convert HTML to PDF
            weasyprint.HTML(filename=html_path).write_pdf(pdf_path)
            
            return pdf_path
            
        except Exception as e:
            self.logger.warning(f"PDF generation failed: {e}")
            return ""
    
    async def _generate_docx_report(self, query_data: Dict, results: List[Dict], summary: str, report_id: str) -> str:
        """Generate Word document report"""
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('CCTNS Database Analysis Report', 0)
            title.alignment = 1  # Center alignment
            
            # Metadata
            doc.add_heading('Report Information', level=1)
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('Report ID', report_id),
                ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                ('Query', query_data.get('original_query', 'N/A')),
                ('Results Count', str(len(results)))
            ]
            
            for i, (key, value) in enumerate(info_data):
                info_table.cell(i, 0).text = key
                info_table.cell(i, 1).text = value
            
            # Summary
            doc.add_heading('AI-Generated Summary', level=1)
            doc.add_paragraph(summary)
            
            # Results (first 50 records)
            if results:
                doc.add_heading('Query Results', level=1)
                doc.add_paragraph(f'Showing first 50 of {len(results)} records:')
                
                # Create table
                if len(results) > 0:
                    headers = list(results[0].keys())
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.style = 'Table Grid'
                    
                    # Add headers
                    for i, header in enumerate(headers):
                        table.cell(0, i).text = header.replace('_', ' ').title()
                    
                    # Add data (limit to 50 rows)
                    for row_data in results[:50]:
                        row_cells = table.add_row().cells
                        for i, header in enumerate(headers):
                            value = row_data.get(header, '')
                            row_cells[i].text = str(value) if value is not None else 'N/A'
            
            docx_path = str(self.reports_dir / f"{report_id}.docx")
            doc.save(docx_path)
            
            return docx_path
            
        except Exception as e:
            self.logger.warning(f"DOCX generation failed: {e}")
            return ""
    
    async def _generate_excel_report(self, results: List[Dict], report_id: str) -> str:
        """Generate Excel report with data and charts"""
        try:
            if not results:
                return ""
            
            excel_path = str(self.reports_dir / f"{report_id}.xlsx")
            
            # Create DataFrame and save to Excel
            df = pd.DataFrame(results)
            
            with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                # Data sheet
                df.to_excel(writer, sheet_name='Data', index=False)
                
                # Summary sheet
                summary_data = {
                    'Metric': ['Total Records', 'Columns', 'Numeric Columns', 'Text Columns'],
                    'Value': [
                        len(df),
                        len(df.columns),
                        len(df.select_dtypes(include=['number']).columns),
                        len(df.select_dtypes(include=['object']).columns)
                    ]
                }
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='Summary', index=False)
                
                # Statistics sheet (if numeric data exists)
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    stats_df = df[numeric_cols].describe()
                    stats_df.to_excel(writer, sheet_name='Statistics')
            
            return excel_path
            
        except Exception as e:
            self.logger.warning(f"Excel generation failed: {e}")
            return ""
    
    def get_report_status(self, report_id: str) -> Dict[str, Any]:
        """Get status of generated reports"""
        report_files = {
            'html': self.reports_dir / f"{report_id}.html",
            'pdf': self.reports_dir / f"{report_id}.pdf",
            'docx': self.reports_dir / f"{report_id}.docx",
            'excel': self.reports_dir / f"{report_id}.xlsx"
        }
        
        status = {}
        for format_name, file_path in report_files.items():
            status[format_name] = {
                'exists': file_path.exists(),
                'path': str(file_path) if file_path.exists() else None,
                'size': file_path.stat().st_size if file_path.exists() else 0
            }
        
        return status