"""
Reports API endpoints for generating and managing reports
"""
from fastapi import APIRouter, HTTPException, Depends, Query, BackgroundTasks
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field
from typing import Dict, Any, List, Optional, Union
import json
import asyncio
import logging
from datetime import datetime, date, timedelta
from pathlib import Path
import uuid
import io

from agents.execution_agent import ExecutionAgent
from agents.visualization_agent import VisualizationAgent
from api.middleware.auth import AuthMiddleware
from config.settings import settings

# Import report generation libraries
try:
    import pandas as pd
    import matplotlib.pyplot as plt
    import seaborn as sns
    from jinja2 import Template
    import pdfkit
    from docx import Document
    from docx.shared import Inches
    REPORT_LIBS_AVAILABLE = True
except ImportError:
    REPORT_LIBS_AVAILABLE = False

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/reports", tags=["reports"])

# Authentication middleware
auth = AuthMiddleware()

# Global agents
execution_agent = None
visualization_agent = None

# Pydantic models
class ReportRequest(BaseModel):
    title: str = Field(..., description="Report title")
    report_type: str = Field(..., description="Type of report: summary, detailed, analytical, dashboard")
    query: Optional[str] = Field(None, description="SQL query for data")
    filters: Optional[Dict[str, Any]] = Field({}, description="Data filters")
    date_range: Optional[Dict[str, str]] = Field(None, description="Date range filter")
    districts: Optional[List[str]] = Field(None, description="District filter")
    crime_types: Optional[List[str]] = Field(None, description="Crime type filter")
    format: str = Field("pdf", description="Output format: pdf, html, docx, xlsx, json")
    include_charts: bool = Field(True, description="Include visualizations")
    template: Optional[str] = Field("default", description="Report template")

class QuickReportRequest(BaseModel):
    report_name: str = Field(..., description="Predefined report name")
    parameters: Optional[Dict[str, Any]] = Field({}, description="Report parameters")
    date_range: Optional[Dict[str, str]] = Field(None, description="Date range")
    format: str = Field("pdf", description="Output format")

class ReportResponse(BaseModel):
    success: bool
    report_id: str
    title: str
    format: str
    file_path: Optional[str] = None
    download_url: Optional[str] = None
    data_summary: Optional[Dict[str, Any]] = None
    generation_time: float
    created_at: str

class ReportStatus(BaseModel):
    report_id: str
    status: str  # pending, processing, completed, failed
    progress: float  # 0-100
    message: str
    created_at: str
    completed_at: Optional[str] = None

class ScheduledReportRequest(BaseModel):
    title: str
    report_config: ReportRequest
    schedule: str  # cron expression
    recipients: List[str]  # email addresses
    enabled: bool = True

# Global report storage
report_cache: Dict[str, Dict] = {}
scheduled_reports: Dict[str, Dict] = {}

@router.on_event("startup")
async def startup_reports_service():
    """Initialize reports service"""
    global execution_agent, visualization_agent
    
    try:
        # Initialize agents
        config = {}  # Would load from config
        execution_agent = ExecutionAgent(config)
        visualization_agent = VisualizationAgent(config)
        
        # Ensure reports directory exists
        reports_dir = Path(settings.REPORTS_DIR)
        reports_dir.mkdir(exist_ok=True)
        
        logger.info("📊 Reports service initialized successfully")
        
    except Exception as e:
        logger.error(f"❌ Reports service initialization failed: {e}")
        raise

@router.post("/generate", response_model=ReportResponse)
async def generate_report(
    request: ReportRequest,
    background_tasks: BackgroundTasks,
    current_user: Dict = Depends(auth.get_current_user)
) -> ReportResponse:
    """Generate a custom report"""
    
    start_time = datetime.now()
    report_id = f"report_{uuid.uuid4().hex[:8]}"
    
    try:
        # Validate request
        if not REPORT_LIBS_AVAILABLE:
            raise HTTPException(
                status_code=500, 
                detail="Report generation libraries not installed"
            )
        
        # Initialize report status
        report_cache[report_id] = {
            "status": "processing",
            "progress": 0,
            "message": "Starting report generation",
            "created_at": start_time.isoformat(),
            "user_id": current_user.get("user_id"),
            "request": request.dict()
        }
        
        # Generate report in background
        background_tasks.add_task(
            _generate_report_background,
            report_id,
            request,
            current_user
        )
        
        return ReportResponse(
            success=True,
            report_id=report_id,
            title=request.title,
            format=request.format,
            generation_time=0.0,
            created_at=start_time.isoformat()
        )
        
    except Exception as e:
        logger.error(f"❌ Report generation request failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/quick", response_model=ReportResponse)
async def generate_quick_report(
    request: QuickReportRequest,
    background_tasks: BackgroundTasks,
    current_user: Dict = Depends(auth.get_current_user)
) -> ReportResponse:
    """Generate a predefined quick report"""
    
    try:
        # Map quick report to full report request
        report_config = await _get_quick_report_config(request.report_name, request.parameters)
        
        full_request = ReportRequest(
            title=f"Quick Report: {request.report_name}",
            report_type="summary",
            format=request.format,
            **report_config
        )
        
        if request.date_range:
            full_request.date_range = request.date_range
        
        return await generate_report(full_request, background_tasks, current_user)
        
    except Exception as e:
        logger.error(f"❌ Quick report generation failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/status/{report_id}", response_model=ReportStatus)
async def get_report_status(
    report_id: str,
    current_user: Dict = Depends(auth.get_current_user)
) -> ReportStatus:
    """Get report generation status"""
    
    if report_id not in report_cache:
        raise HTTPException(status_code=404, detail="Report not found")
    
    report_data = report_cache[report_id]
    
    # Check user authorization
    if report_data.get("user_id") != current_user.get("user_id") and current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Access denied")
    
    return ReportStatus(
        report_id=report_id,
        status=report_data.get("status", "unknown"),
        progress=report_data.get("progress", 0),
        message=report_data.get("message", ""),
        created_at=report_data.get("created_at", ""),
        completed_at=report_data.get("completed_at")
    )

@router.get("/download/{report_id}")
async def download_report(
    report_id: str,
    current_user: Dict = Depends(auth.get_current_user)
):
    """Download generated report"""
    
    if report_id not in report_cache:
        raise HTTPException(status_code=404, detail="Report not found")
    
    report_data = report_cache[report_id]
    
    # Check user authorization
    if report_data.get("user_id") != current_user.get("user_id") and current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Access denied")
    
    if report_data.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Report not ready for download")
    
    file_path = report_data.get("file_path")
    if not file_path or not Path(file_path).exists():
        raise HTTPException(status_code=404, detail="Report file not found")
    
    # Determine media type
    format = report_data.get("format", "pdf")
    media_types = {
        "pdf": "application/pdf",
        "html": "text/html",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "json": "application/json",
        "csv": "text/csv"
    }
    
    media_type = media_types.get(format, "application/octet-stream")
    filename = f"{report_data.get('title', 'report')}_{report_id}.{format}"
    
    return FileResponse(
        path=file_path,
        media_type=media_type,
        filename=filename
    )

@router.get("/list")
async def list_reports(
    status: Optional[str] = Query(None, description="Filter by status"),
    limit: int = Query(50, description="Number of reports to return"),
    offset: int = Query(0, description="Offset for pagination"),
    current_user: Dict = Depends(auth.get_current_user)
):
    """List user's reports"""
    
    user_id = current_user.get("user_id")
    is_admin = current_user.get("role") == "admin"
    
    # Filter reports
    filtered_reports = []
    for report_id, report_data in report_cache.items():
        # Check access
        if not is_admin and report_data.get("user_id") != user_id:
            continue
        
        # Filter by status
        if status and report_data.get("status") != status:
            continue
        
        filtered_reports.append({
            "report_id": report_id,
            "title": report_data.get("request", {}).get("title", "Unknown"),
            "status": report_data.get("status"),
            "format": report_data.get("request", {}).get("format"),
            "created_at": report_data.get("created_at"),
            "completed_at": report_data.get("completed_at")
        })
    
    # Sort by creation date (newest first)
    filtered_reports.sort(key=lambda x: x["created_at"], reverse=True)
    
    # Paginate
    total = len(filtered_reports)
    paginated_reports = filtered_reports[offset:offset + limit]
    
    return {
        "reports": paginated_reports,
        "total": total,
        "limit": limit,
        "offset": offset
    }

@router.delete("/delete/{report_id}")
async def delete_report(
    report_id: str,
    current_user: Dict = Depends(auth.get_current_user)
):
    """Delete a report"""
    
    if report_id not in report_cache:
        raise HTTPException(status_code=404, detail="Report not found")
    
    report_data = report_cache[report_id]
    
    # Check user authorization
    if report_data.get("user_id") != current_user.get("user_id") and current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Access denied")
    
    # Delete file if exists
    file_path = report_data.get("file_path")
    if file_path and Path(file_path).exists():
        try:
            Path(file_path).unlink()
        except Exception as e:
            logger.warning(f"Failed to delete report file {file_path}: {e}")
    
    # Remove from cache
    del report_cache[report_id]
    
    return {"message": "Report deleted successfully"}

@router.get("/templates")
async def get_report_templates(current_user: Dict = Depends(auth.get_current_user)):
    """Get available report templates"""
    
    templates = {
        "default": {
            "name": "Default Report",
            "description": "Standard police report template",
            "sections": ["summary", "data", "charts", "conclusion"]
        },
        "crime_summary": {
            "name": "Crime Summary",
            "description": "Summary of crime statistics",
            "sections": ["overview", "crime_types", "district_breakdown", "trends"]
        },
        "fir_analysis": {
            "name": "FIR Analysis",
            "description": "Detailed FIR analysis report",
            "sections": ["fir_summary", "status_breakdown", "officer_performance", "resolution_rates"]
        },
        "district_report": {
            "name": "District Report",
            "description": "Comprehensive district-wise analysis",
            "sections": ["district_overview", "crime_statistics", "station_performance", "resource_allocation"]
        },
        "monthly_summary": {
            "name": "Monthly Summary",
            "description": "Monthly crime and police activity summary",
            "sections": ["monthly_overview", "crime_trends", "arrest_statistics", "performance_metrics"]
        }
    }
    
    return {"templates": templates}

@router.get("/quick-reports")
async def get_quick_reports(current_user: Dict = Depends(auth.get_current_user)):
    """Get available quick reports"""
    
    quick_reports = {
        "daily_summary": {
            "name": "Daily Summary",
            "description": "Daily police activity summary",
            "parameters": ["date"]
        },
        "crime_stats": {
            "name": "Crime Statistics",
            "description": "Crime statistics for selected period",
            "parameters": ["start_date", "end_date", "district"]
        },
        "fir_status": {
            "name": "FIR Status Report",
            "description": "Status of FIRs in the system",
            "parameters": ["status", "district", "date_range"]
        },
        "officer_performance": {
            "name": "Officer Performance",
            "description": "Performance metrics for officers",
            "parameters": ["district", "month", "year"]
        },
        "district_comparison": {
            "name": "District Comparison",
            "description": "Compare crime statistics across districts",
            "parameters": ["districts", "crime_types", "time_period"]
        }
    }
    
    return {"quick_reports": quick_reports}

@router.post("/schedule", response_model=Dict[str, str])
async def schedule_report(
    request: ScheduledReportRequest,
    current_user: Dict = Depends(auth.get_current_user)
):
    """Schedule a recurring report"""
    
    schedule_id = f"schedule_{uuid.uuid4().hex[:8]}"
    
    scheduled_reports[schedule_id] = {
        "id": schedule_id,
        "title": request.title,
        "config": request.report_config.dict(),
        "schedule": request.schedule,
        "recipients": request.recipients,
        "enabled": request.enabled,
        "created_by": current_user.get("user_id"),
        "created_at": datetime.now().isoformat(),
        "last_run": None,
        "next_run": None  # Would calculate based on cron
    }
    
    logger.info(f"📅 Scheduled report created: {schedule_id}")
    
    return {
        "message": "Report scheduled successfully",
        "schedule_id": schedule_id
    }

@router.get("/scheduled")
async def list_scheduled_reports(current_user: Dict = Depends(auth.get_current_user)):
    """List scheduled reports"""
    
    user_id = current_user.get("user_id")
    is_admin = current_user.get("role") == "admin"
    
    user_schedules = []
    for schedule_id, schedule_data in scheduled_reports.items():
        if is_admin or schedule_data.get("created_by") == user_id:
            user_schedules.append({
                "schedule_id": schedule_id,
                **schedule_data
            })
    
    return {"scheduled_reports": user_schedules}

@router.delete("/scheduled/{schedule_id}")
async def delete_scheduled_report(
    schedule_id: str,
    current_user: Dict = Depends(auth.get_current_user)
):
    """Delete a scheduled report"""
    
    if schedule_id not in scheduled_reports:
        raise HTTPException(status_code=404, detail="Scheduled report not found")
    
    schedule_data = scheduled_reports[schedule_id]
    
    # Check authorization
    if schedule_data.get("created_by") != current_user.get("user_id") and current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Access denied")
    
    del scheduled_reports[schedule_id]
    
    return {"message": "Scheduled report deleted successfully"}

@router.get("/analytics")
async def get_report_analytics(
    days: int = Query(30, description="Number of days to analyze"),
    current_user: Dict = Depends(auth.get_current_user)
):
    """Get analytics about report usage"""
    
    user_id = current_user.get("user_id")
    is_admin = current_user.get("role") == "admin"
    
    # Calculate date threshold
    threshold_date = datetime.now() - timedelta(days=days)
    
    # Analyze reports
    total_reports = 0
    completed_reports = 0
    failed_reports = 0
    formats_count = {}
    templates_count = {}
    
    for report_data in report_cache.values():
        # Filter by user if not admin
        if not is_admin and report_data.get("user_id") != user_id:
            continue
        
        # Filter by date
        created_at = datetime.fromisoformat(report_data.get("created_at", ""))
        if created_at < threshold_date:
            continue
        
        total_reports += 1
        
        status = report_data.get("status", "unknown")
        if status == "completed":
            completed_reports += 1
        elif status == "failed":
            failed_reports += 1
        
        # Count formats
        format = report_data.get("request", {}).get("format", "unknown")
        formats_count[format] = formats_count.get(format, 0) + 1
        
        # Count templates
        template = report_data.get("request", {}).get("template", "unknown")
        templates_count[template] = templates_count.get(template, 0) + 1
    
    success_rate = (completed_reports / total_reports * 100) if total_reports > 0 else 0
    
    return {
        "period_days": days,
        "total_reports": total_reports,
        "completed_reports": completed_reports,
        "failed_reports": failed_reports,
        "success_rate": round(success_rate, 2),
        "formats_usage": formats_count,
        "templates_usage": templates_count,
        "generated_at": datetime.now().isoformat()
    }

# Background task functions
async def _generate_report_background(
    report_id: str,
    request: ReportRequest,
    current_user: Dict
):
    """Generate report in background task"""
    
    try:
        # Update status
        report_cache[report_id]["progress"] = 10
        report_cache[report_id]["message"] = "Fetching data..."
        
        # Get data based on query or filters
        data = await _fetch_report_data(request)
        
        report_cache[report_id]["progress"] = 40
        report_cache[report_id]["message"] = "Processing data..."
        
        # Generate visualizations if requested
        charts = []
        if request.include_charts and data:
            charts = await _generate_report_charts(data, request.report_type)
        
        report_cache[report_id]["progress"] = 70
        report_cache[report_id]["message"] = "Generating report..."
        
        # Generate report based on format
        file_path = await _generate_report_file(
            report_id,
            request,
            data,
            charts,
            current_user
        )
        
        # Update final status
        report_cache[report_id].update({
            "status": "completed",
            "progress": 100,
            "message": "Report generated successfully",
            "completed_at": datetime.now().isoformat(),
            "file_path": file_path,
            "data_summary": {
                "total_rows": len(data) if data else 0,
                "charts_count": len(charts)
            }
        })
        
        logger.info(f"✅ Report generated successfully: {report_id}")
        
    except Exception as e:
        logger.error(f"❌ Report generation failed for {report_id}: {e}")
        
        report_cache[report_id].update({
            "status": "failed",
            "message": f"Report generation failed: {str(e)}",
            "completed_at": datetime.now().isoformat()
        })

async def _fetch_report_data(request: ReportRequest) -> List[Dict]:
    """Fetch data for report"""
    
    global execution_agent
    
    try:
        # Use provided query or generate based on filters
        if request.query:
            sql_query = request.query
        else:
            sql_query = await _build_query_from_filters(request)
        
        # Execute query
        result = await execution_agent.execute({
            "type": "execute_sql",
            "sql": sql_query,
            "use_cache": True
        })
        
        if result.get("success") and result.get("result", {}).get("success"):
            return result["result"]["data"]
        else:
            logger.error(f"Data fetch failed: {result.get('error', 'Unknown error')}")
            return []
    
    except Exception as e:
        logger.error(f"Data fetch error: {e}")
        return []

async def _build_query_from_filters(request: ReportRequest) -> str:
    """Build SQL query from filters"""
    
    # Base queries for different report types
    base_queries = {
        "summary": "SELECT COUNT(*) as total_records, 'summary' as type FROM FIR",
        "detailed": "SELECT * FROM FIR",
        "analytical": """
            SELECT 
                d.district_name,
                ct.crime_description,
                COUNT(*) as crime_count,
                DATE_TRUNC('month', f.incident_date) as month
            FROM FIR f
            JOIN DISTRICT_MASTER d ON f.district_id = d.district_id
            JOIN CRIME_TYPE_MASTER ct ON f.crime_type_id = ct.crime_type_id
        """,
        "dashboard": """
            SELECT 
                COUNT(*) as total_firs,
                COUNT(CASE WHEN status = 'OPEN' THEN 1 END) as open_firs,
                COUNT(CASE WHEN status = 'CLOSED' THEN 1 END) as closed_firs
            FROM FIR
        """
    }
    
    query = base_queries.get(request.report_type, base_queries["summary"])
    
    # Add filters
    conditions = []
    
    if request.date_range:
        start_date = request.date_range.get("start_date")
        end_date = request.date_range.get("end_date")
        if start_date and end_date:
            conditions.append(f"incident_date BETWEEN '{start_date}' AND '{end_date}'")
    
    if request.districts:
        district_list = "', '".join(request.districts)
        conditions.append(f"d.district_name IN ('{district_list}')")
    
    if request.crime_types:
        crime_list = "', '".join(request.crime_types)
        conditions.append(f"ct.crime_description IN ('{crime_list}')")
    
    # Add WHERE clause if conditions exist
    if conditions:
        if "WHERE" in query.upper():
            query += " AND " + " AND ".join(conditions)
        else:
            query += " WHERE " + " AND ".join(conditions)
    
    # Add GROUP BY for analytical reports
    if request.report_type == "analytical":
        query += " GROUP BY d.district_name, ct.crime_description, DATE_TRUNC('month', f.incident_date)"
        query += " ORDER BY month DESC, crime_count DESC"
    
    query += " LIMIT 1000"  # Safety limit
    
    return query

async def _generate_report_charts(data: List[Dict], report_type: str) -> List[Dict]:
    """Generate charts for report"""
    
    global visualization_agent
    
    try:
        if not data:
            return []
        
        charts = []
        
        # Generate different charts based on report type
        if report_type == "analytical":
            # District-wise crime chart
            chart_result = await visualization_agent.execute({
                "type": "auto_chart",
                "data": data,
                "title": "Crime Distribution by District"
            })
            
            if chart_result.get("success"):
                charts.extend(chart_result["result"]["charts"])
        
        elif report_type == "summary":
            # Simple summary chart
            chart_result = await visualization_agent.execute({
                "type": "specific_chart",
                "data": data,
                "chart_type": "bar",
                "title": "Summary Statistics"
            })
            
            if chart_result.get("success"):
                charts.extend(chart_result["result"]["charts"])
        
        return charts
        
    except Exception as e:
        logger.error(f"Chart generation error: {e}")
        return []

async def _generate_report_file(
    report_id: str,
    request: ReportRequest,
    data: List[Dict],
    charts: List[Dict],
    current_user: Dict
) -> str:
    """Generate report file in requested format"""
    
    reports_dir = Path(settings.REPORTS_DIR)
    file_name = f"{report_id}.{request.format}"
    file_path = reports_dir / file_name
    
    try:
        if request.format == "json":
            await _generate_json_report(file_path, request, data, charts)
        elif request.format == "csv":
            await _generate_csv_report(file_path, data)
        elif request.format == "html":
            await _generate_html_report(file_path, request, data, charts, current_user)
        elif request.format == "pdf":
            await _generate_pdf_report(file_path, request, data, charts, current_user)
        elif request.format == "docx":
            await _generate_docx_report(file_path, request, data, charts, current_user)
        elif request.format == "xlsx":
            await _generate_xlsx_report(file_path, request, data, charts)
        else:
            raise ValueError(f"Unsupported format: {request.format}")
        
        return str(file_path)
        
    except Exception as e:
        logger.error(f"Report file generation failed: {e}")
        raise

async def _generate_json_report(file_path: Path, request: ReportRequest, data: List[Dict], charts: List[Dict]):
    """Generate JSON report"""
    
    report_content = {
        "title": request.title,
        "report_type": request.report_type,
        "generated_at": datetime.now().isoformat(),
        "data": data,
        "charts": charts,
        "summary": {
            "total_records": len(data),
            "charts_count": len(charts)
        }
    }
    
    with open(file_path, 'w') as f:
        json.dump(report_content, f, indent=2, default=str)

async def _generate_csv_report(file_path: Path, data: List[Dict]):
    """Generate CSV report"""
    
    if data:
        df = pd.DataFrame(data)
        df.to_csv(file_path, index=False)
    else:
        # Create empty CSV
        with open(file_path, 'w') as f:
            f.write("No data available\n")

async def _generate_html_report(file_path: Path, request: ReportRequest, data: List[Dict], charts: List[Dict], current_user: Dict):
    """Generate HTML report"""
    
    # HTML template
    html_template = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>{{ title }}</title>
        <meta charset="utf-8">
        <style>
            body { font-family: Arial, sans-serif; margin: 40px; }
            .header { text-align: center; margin-bottom: 30px; }
            .summary { background: #f5f5f5; padding: 20px; margin: 20px 0; }
            .data-table { width: 100%; border-collapse: collapse; margin: 20px 0; }
            .data-table th, .data-table td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            .data-table th { background-color: #f2f2f2; }
            .chart { margin: 30px 0; text-align: center; }
            .footer { margin-top: 50px; text-align: center; color: #666; }
        </style>
    </head>
    <body>
        <div class="header">
            <h1>{{ title }}</h1>
            <p>Report Type: {{ report_type }}</p>
            <p>Generated on: {{ generated_at }}</p>
            <p>Generated by: {{ user_name }}</p>
        </div>
        
        <div class="summary">
            <h2>Summary</h2>
            <p>Total Records: {{ total_records }}</p>
            <p>Charts: {{ charts_count }}</p>
        </div>
        
        {% if data %}
        <div class="data-section">
            <h2>Data</h2>
            <table class="data-table">
                <thead>
                    <tr>
                        {% for column in columns %}
                        <th>{{ column }}</th>
                        {% endfor %}
                    </tr>
                </thead>
                <tbody>
                    {% for row in data[:100] %}
                    <tr>
                        {% for value in row.values() %}
                        <td>{{ value }}</td>
                        {% endfor %}
                    </tr>
                    {% endfor %}
                </tbody>
            </table>
            {% if data|length > 100 %}
            <p><em>Showing first 100 records out of {{ data|length }}</em></p>
            {% endif %}
        </div>
        {% endif %}
        
        {% for chart in charts %}
        {% if chart.format == 'html' %}
        <div class="chart">
            <h3>{{ chart.title }}</h3>
            {{ chart.data|safe }}
        </div>
        {% endif %}
        {% endfor %}
        
        <div class="footer">
            <p>Generated by CCTNS Copilot Engine</p>
        </div>
    </body>
    </html>
    """
    
    template = Template(html_template)
    
    columns = list(data[0].keys()) if data else []
    
    html_content = template.render(
        title=request.title,
        report_type=request.report_type,
        generated_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        user_name=current_user.get("username", "Unknown"),
        total_records=len(data),
        charts_count=len(charts),
        data=data,
        columns=columns,
        charts=charts
    )
    
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

async def _generate_pdf_report(file_path: Path, request: ReportRequest, data: List[Dict], charts: List[Dict], current_user: Dict):
    """Generate PDF report"""
    
    # First generate HTML
    html_path = file_path.with_suffix('.html')
    await _generate_html_report(html_path, request, data, charts, current_user)
    
    # Convert HTML to PDF
    try:
        pdfkit.from_file(str(html_path), str(file_path))
        # Clean up temporary HTML file
        html_path.unlink()
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        # Fallback: just copy HTML file with PDF extension
        import shutil
        shutil.move(str(html_path), str(file_path))

async def _generate_docx_report(file_path: Path, request: ReportRequest, data: List[Dict], charts: List[Dict], current_user: Dict):
    """Generate DOCX report"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading(request.title, 0)
    title.alignment = 1  # Center alignment
    
    # Metadata
    doc.add_paragraph(f"Report Type: {request.report_type}")
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Generated by: {current_user.get('username', 'Unknown')}")
    
    # Summary
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f"Total Records: {len(data)}")
    doc.add_paragraph(f"Charts: {len(charts)}")
    
    # Data table (first 50 rows)
    if data:
        doc.add_heading('Data', level=1)
        
        # Create table
        columns = list(data[0].keys())
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(columns):
            hdr_cells[i].text = str(column)
        
        # Data rows (limit to 50)
        for row_data in data[:50]:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data.values()):
                row_cells[i].text = str(value)
        
        if len(data) > 50:
            doc.add_paragraph(f"Showing first 50 records out of {len(data)}")
    
    # Save document
    doc.save(str(file_path))

async def _generate_xlsx_report(file_path: Path, request: ReportRequest, data: List[Dict], charts: List[Dict]):
    """Generate XLSX report"""
    
    with pd.ExcelWriter(str(file_path), engine='openpyxl') as writer:
        # Data sheet
        if data:
            df = pd.DataFrame(data)
            df.to_excel(writer, sheet_name='Data', index=False)
        
        # Summary sheet
        summary_data = {
            'Metric': ['Total Records', 'Charts Generated', 'Report Type', 'Generated At'],
            'Value': [
                len(data),
                len(charts),
                request.report_type,
                datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            ]
        }
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, sheet_name='Summary', index=False)

async def _get_quick_report_config(report_name: str, parameters: Dict[str, Any]) -> Dict[str, Any]:
    """Get configuration for quick report"""
    
    configs = {
        "daily_summary": {
            "query": """
                SELECT 
                    COUNT(*) as total_firs,
                    COUNT(CASE WHEN status = 'OPEN' THEN 1 END) as open_firs,
                    COUNT(CASE WHEN status = 'CLOSED' THEN 1 END) as closed_firs
                FROM FIR 
                WHERE DATE(incident_date) = CURRENT_DATE
            """,
            "report_type": "summary",
            "include_charts": True
        },
        "crime_stats": {
            "query": """
                SELECT 
                    ct.crime_description,
                    COUNT(*) as count,
                    d.district_name
                FROM FIR f
                JOIN CRIME_TYPE_MASTER ct ON f.crime_type_id = ct.crime_type_id
                JOIN DISTRICT_MASTER d ON f.district_id = d.district_id
                GROUP BY ct.crime_description, d.district_name
                ORDER BY count DESC
            """,
            "report_type": "analytical",
            "include_charts": True
        },
        "fir_status": {
            "query": """
                SELECT 
                    status,
                    COUNT(*) as count,
                    AVG(EXTRACT(DAY FROM (CURRENT_DATE - incident_date))) as avg_age_days
                FROM FIR
                GROUP BY status
            """,
            "report_type": "summary",
            "include_charts": True
        },
        "officer_performance": {
            "query": """
                SELECT 
                    o.officer_name,
                    o.rank,
                    d.district_name,
                    COUNT(a.arrest_id) as total_arrests,
                    COUNT(DISTINCT f.fir_id) as cases_handled
                FROM OFFICER_MASTER o
                LEFT JOIN ARREST a ON o.officer_id = a.officer_id
                LEFT JOIN FIR f ON a.fir_id = f.fir_id
                LEFT JOIN STATION_MASTER s ON o.station_id = s.station_id
                LEFT JOIN DISTRICT_MASTER d ON s.district_id = d.district_id
                GROUP BY o.officer_name, o.rank, d.district_name
                ORDER BY total_arrests DESC
            """,
            "report_type": "analytical",
            "include_charts": True
        },
        "district_comparison": {
            "query": """
                SELECT 
                    d.district_name,
                    COUNT(f.fir_id) as total_firs,
                    COUNT(a.arrest_id) as total_arrests,
                    COUNT(CASE WHEN f.status = 'CLOSED' THEN 1 END) as closed_cases,
                    ROUND(COUNT(CASE WHEN f.status = 'CLOSED' THEN 1 END) * 100.0 / COUNT(f.fir_id), 2) as closure_rate
                FROM DISTRICT_MASTER d
                LEFT JOIN FIR f ON d.district_id = f.district_id
                LEFT JOIN ARREST a ON f.fir_id = a.fir_id
                GROUP BY d.district_name
                ORDER BY total_firs DESC
            """,
            "report_type": "analytical",
            "include_charts": True
        }
    }
    
    config = configs.get(report_name, {})
    
    # Apply parameters to modify query if needed
    query = config.get("query", "")
    
    # Date range parameters
    if "start_date" in parameters and "end_date" in parameters:
        start_date = parameters["start_date"]
        end_date = parameters["end_date"]
        if "WHERE" in query.upper():
            query += f" AND f.incident_date BETWEEN '{start_date}' AND '{end_date}'"
        else:
            query += f" WHERE f.incident_date BETWEEN '{start_date}' AND '{end_date}'"
    
    # District filter
    if "district" in parameters:
        district = parameters["district"]
        if "WHERE" in query.upper():
            query += f" AND d.district_name = '{district}'"
        else:
            query += f" WHERE d.district_name = '{district}'"
    
    # Status filter
    if "status" in parameters:
        status = parameters["status"]
        if "WHERE" in query.upper():
            query += f" AND f.status = '{status}'"
        else:
            query += f" WHERE f.status = '{status}'"
    
    # Update the query in config
    if query != config.get("query", ""):
        config["query"] = query
    
    return config